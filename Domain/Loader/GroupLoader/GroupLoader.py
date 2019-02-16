import re
from collections import namedtuple
from typing import List, Type

from docx import Document

from DataBase2 import Student, Group
from Domain.Loader import WordReader, Reader


class GroupLoader(Reader):
    def get_group(self) -> Group:
        raise NotImplementedError()

    def get_students_list(self) -> List[Student]:
        raise NotImplementedError()


class GroupWordLoader(WordReader, GroupLoader):
    def get_students_list(self):
        return self.students

    def get_group(self) -> Group:
        return self.group

    group_regex = re.compile(
        r'(?:(?:[Гг]руппа)\s)?([А-Яа-я]+\s?-\s?[0-9]+)(?:\s(?:[Гг]руппа))?'
    )

    full_name_regex = re.compile(
        r'^\s*((?:(?:[А-Яа-я])+(?:-[А-Яа-я]+)?\s){1,2}(?:(?:[А-Яа-я])+(?:-[А-Яа-я]+)?))\s*(?:[,.;])?\s*(?:[,.;])?$'
    )

    start_table = namedtuple('start_table', 'table row col')

    def __init__(self, file, professor, session):
        self.doc: Document = Document(str(file))

        self.mode: GroupWordLoader.Mode = None
        self.start_paragraph: int = None

        self.professor = professor
        self.session = session

        self._get_group()
        self._get_students_list()

    def _get_group(self):
        for row_index, paragraph in enumerate(self.doc.paragraphs):
            groups = self.group_regex.findall(paragraph.text)

            if len(groups) > 0 and groups[0] is not None and groups[0] != '':
                self.mode = WordReader.Mode.LIST
                self.start_paragraph = row_index + 1

                self.group_name = groups[0].replace(' ', '')

        for table in self.doc.tables:
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(table.row_cells(row_index)):
                    groups = self.group_regex.findall(cell.text)

                    if len(groups) > 0 and groups[0] is not None and groups[0] != '':
                        self.mode = WordReader.Mode.LIST_IN_TABLE
                        self.start_table = GroupWordLoader.start_table(table, row_index + 1, col_index)

                        self.group_name = groups[0].replace(' ', '')

        if self.group_name not in [None, '']:
            self.group = Group.get_or_create(self.session, name=self.group_name)
        else:
            raise ValueError('group name is None')

    def _get_students_list(self):
        def get_student(text):
            full_name = self.full_name_regex.findall(text)
            if full_name is not None and len(full_name) > 0 and full_name[0] != '':
                splited = full_name[0].split(' ')
                student = Student.get_or_create(
                    last_name=splited[0],
                    first_name=splited[1],
                    middle_name=splited[2] if len(splited) >= 3 else ''
                )
                if self.group not in student.groups:
                    student.groups.append(self.group)
                return student

        students = []
        if self.mode == WordReader.Mode.LIST:
            for paragraph in self.doc.paragraphs[self.start_paragraph:]:
                student = get_student(paragraph.text)
                if student is not None:
                    students.append(student)

            if len(students) == 0:
                for table in self.doc.tables:
                    for cell in table._cells:
                        student = get_student(cell.text)
                        if student is not None:
                            students.append(student)

        elif self.mode == WordReader.Mode.LIST_IN_TABLE:
            start_row = self.start_table.row
            main_col = self.start_table.col
            table = self.start_table.table
            for row_index, row in enumerate(table.rows[start_row:], start_row):
                cells = table.row_cells(row_index)
                if len(cells) >= main_col:
                    cell = cells[main_col]
                    student = get_student(cell.text)
                    if student is not None:
                        students.append(student)

        self.students = sorted(students, key=lambda x: str(x))
